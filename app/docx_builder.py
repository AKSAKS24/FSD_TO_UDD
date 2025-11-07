import io
from typing import List, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

def _set_heading_run_style(run, size=16, bold=True, color="2F5496"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def _add_footer_with_page_numbers(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1); r_element.append(instrText); r_element.append(fldChar2)
    p.add_run(" of ")
    run2 = p.add_run()
    fldChar1b = OxmlElement('w:fldChar'); fldChar1b.set(qn('w:fldCharType'), 'begin')
    instrTextb = OxmlElement('w:instrText'); instrTextb.set(qn('xml:space'), 'preserve'); instrTextb.text = 'NUMPAGES'
    fldChar2b = OxmlElement('w:fldChar'); fldChar2b.set(qn('w:fldCharType'), 'end')
    r2 = run2._r
    r2.append(fldChar1b); r2.append(instrTextb); r2.append(fldChar2b)

def _add_toc(doc: Document, title="Table of Contents"):
    p = doc.add_paragraph()
    run = p.add_run(title)
    _set_heading_run_style(run, size=16, bold=True, color="1F4E79")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = doc.add_paragraph()
    fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'TOC \\o \"1-3\" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    r = p2.add_run()._r
    r.append(fldChar); r.append(instrText); r.append(fldChar2); r.append(fldChar3)
    doc.add_page_break()

def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    color = "2F5496" if level == 1 else "44546A"
    size = 16 if level == 1 else 13
    _set_heading_run_style(run, size=size, bold=True, color=color)
    try:
        p.style = doc.styles[f'Heading {level}']
    except KeyError:
        pass

def _add_markdown_table(doc: Document, md: str):
    lines = [l.strip() for l in md.strip().splitlines() if l.strip()]
    table_lines = [l for l in lines if l.startswith("|") and l.endswith("|")]
    if not table_lines:
        doc.add_paragraph(md); return
    clean_rows = []
    for row in table_lines:
        cells = [c.strip() for c in row.strip("|").split("|")]
        clean_rows.append(cells)
    if len(clean_rows) < 2:
        doc.add_paragraph(md); return
    header = clean_rows[0]
    data_rows = [r for r in clean_rows[1:] if not all(set(c) <= {"-", ":"} for c in r)]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for j, h in enumerate(header):
        hdr_cells[j].text = h
    for r in data_rows:
        row_cells = t.add_row().cells
        for j, val in enumerate(r):
            if j < len(row_cells):
                row_cells[j].text = val

def build_docx(udd_sections: List[Tuple[str, str]], title: str = "Unified Design Document") -> bytes:
    doc = Document()

    tp = doc.add_paragraph()
    tr = tp.add_run(title)
    _set_heading_run_style(tr, size=20, bold=True, color="1F4E79")
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_toc(doc)

    for name, content in udd_sections:
        _add_heading(doc, name, level=1)
        chunks = [seg for seg in content.split("\n\n") if seg.strip()]
        for seg in chunks:
            if seg.strip().startswith("|") and "|" in seg.strip():
                _add_markdown_table(doc, seg)
            else:
                para = doc.add_paragraph(seg)
                para.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()

    _add_footer_with_page_numbers(doc)

    buff = io.BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff.read()
